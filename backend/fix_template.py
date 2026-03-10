from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# Page Size aur Margins
for section in doc.sections:
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(12)

# Dynamic Header
academy_para = doc.add_paragraph()
academy_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
academy_run = academy_para.add_run('{{ academy_name }}')
academy_run.bold = True
academy_run.font.size = Pt(18)

syllabus_para = doc.add_paragraph()
syllabus_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
syllabus_run = syllabus_para.add_run('Syllabus: {{ syllabus }}')
syllabus_run.bold = True
syllabus_run.font.size = Pt(14)

doc.add_paragraph('Name: ____________________        Date:   {{ date }}         Time:   {{ time }}          Max Marks:  {{ total_marks }}')
doc.add_paragraph('Name: ____________________        Date:   {{ date }}         Time:   {{ time }}          Max Marks:  {{ total_marks }}')

# MCQs
doc.add_paragraph()
mcq_heading = doc.add_paragraph()
mcq_run = mcq_heading.add_run('Multiple Choice Questions (MCQs)')
mcq_run.bold = True
doc.add_paragraph('Choose the correct one from the following options \t\t\t\t/{{ mcq_total }}')

# Table Setup
table = doc.add_table(rows=2, cols=6)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False

widths = [Inches(0.4), Inches(3.4), Inches(0.8), Inches(0.8), Inches(0.8), Inches(0.8)]
for row in table.rows:
    for idx, width in enumerate(widths):
        row.cells[idx].width = width

headers = ['No.', 'Question', 'A', 'B', 'C', 'D']
hdr_cells = table.rows[0].cells
for i, text in enumerate(headers):
    run = hdr_cells[i].paragraphs[0].add_run(text)
    run.bold = True

# ---- CLEAN TAG SEPARATION (The Magic Fix) ----
row_cells = table.rows[1].cells

# Pehla cell: Tag alag line me, number alag line me
row_cells[0].text = '{% tr for m in mcqs %}'
row_cells[0].add_paragraph('{{ loop.index }}')

row_cells[1].text = '{{ m.question }}'
row_cells[2].text = '{{ m.a }}'
row_cells[3].text = '{{ m.b }}'
row_cells[4].text = '{{ m.c }}'

# Aakhri cell: Text alag line me, End tag alag line me
row_cells[5].text = '{{ m.d }}'
row_cells[5].add_paragraph('{% tr endfor %}')

# Short Questions
doc.add_paragraph()
short_heading = doc.add_paragraph()
short_run = short_heading.add_run('Short Questions (2 marks each) \t\t\t\t\t\t/{{ short_total }}')
short_run.bold = True
doc.add_paragraph('{% p for sq in short_qs %}')
doc.add_paragraph('{{ loop.index }}. {{ sq.text }}')
doc.add_paragraph('{% p endfor %}')

# Long Questions
doc.add_paragraph()
long_heading = doc.add_paragraph()
long_run = long_heading.add_run('Long Question ({{ long_instruction }} marks) \t\t\t\t\t/{{ long_total }}')
long_run.bold = True
doc.add_paragraph('{% p for lq in long_qs %}')
doc.add_paragraph('{{ loop.index }}. {{ lq.text }}')
doc.add_paragraph('{% p endfor %}')

# Save the file
current_dir = os.path.dirname(os.path.abspath(__file__))
templates_dir = os.path.join(current_dir, "templates")
os.makedirs(templates_dir, exist_ok=True)
file_path = os.path.join(templates_dir, 'ultimate_template.docx')

doc.save(file_path)
print(f"Success! Python-friendly template saved here: {file_path}")