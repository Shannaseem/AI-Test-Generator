import os
import re
import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.oxml.shared import OxmlElement, qn

def clean_text(text):
    return re.sub(r'^\d+[\.\)]\s*', '', str(text)).strip()

def set_rtl_paragraph(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def add_urdu_run(paragraph, text, font_size=12, bold=False):
    run = paragraph.add_run(text)
    run.font.name = 'Jameel Noori Nastaleeq'
    run.font.size = Pt(font_size)
    run.bold = bold
    rPr = run._r.get_or_add_rPr()
    rtl = OxmlElement('w:rtl')
    rPr.append(rtl)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:cs'), 'Jameel Noori Nastaleeq')
    return run

def remove_table_borders(table):
    tbl = table._tbl
    tbl_pr_list = tbl.xpath('./w:tblPr')
    if tbl_pr_list:
        tbl_pr = tbl_pr_list[0]
    else:
        tbl_pr = OxmlElement('w:tblPr')
        tbl.insert(0, tbl_pr)
        
    tbl_borders_list = tbl_pr.xpath('./w:tblBorders')
    if tbl_borders_list:
        tbl_borders = tbl_borders_list[0]
        tbl_borders.clear() 
    else:
        tbl_borders = OxmlElement('w:tblBorders')
        tbl_pr.append(tbl_borders)

    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tbl_borders.append(border)

# THE PERFECTED INVISIBLE SIDE-BY-SIDE FUNCTION
def add_side_by_side_bilingual(parent, eng_text, urdu_text, num_prefix="", font_size=12, space_after=6, is_bold=False):
    table = parent.add_table(rows=1, cols=2)
    table.autofit = False
    
    # Force exactly 50% width for each side to push Urdu to the far right
    table.columns[0].width = Inches(3.83)
    table.columns[1].width = Inches(3.83)
    
    remove_table_borders(table)

    row = table.rows[0]
    
    # English Side (Strictly Left)
    eng_cell = row.cells[0]
    eng_cell.width = Inches(3.83)
    eng_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    p_eng = eng_cell.paragraphs[0]
    p_eng.paragraph_format.space_after = Pt(space_after)
    p_eng.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if num_prefix:
        r_num = p_eng.add_run(num_prefix)
        r_num.bold = True
        r_num.font.size = Pt(font_size)

    r_eng = p_eng.add_run(eng_text)
    r_eng.font.size = Pt(font_size)
    if is_bold: r_eng.bold = True

    # Urdu Side (Strictly Right aligned Nastaleeq)
    urdu_cell = row.cells[1]
    urdu_cell.width = Inches(3.83)
    urdu_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    p_urdu = urdu_cell.paragraphs[0]
    p_urdu.paragraph_format.space_after = Pt(space_after)
    set_rtl_paragraph(p_urdu)
    add_urdu_run(p_urdu, urdu_text, font_size=font_size + 2, bold=is_bold)

    return table


# ==========================================
# MAIN FUNCTION
# ==========================================
def generate_word_file(academy_name, subject, class_name, test_date, time_allowed,
                       syllabus, long_q_marks, ai_data, template_style,
                       short_total="8", short_attempt="5", exam_pattern="chapter",
                       short_groups="1", long_total="3", long_attempt="2", bilingual="no",
                       generate_answer_key="no"):

    doc = Document()
    is_bilingual = (bilingual == "yes")

    if exam_pattern == "board":
        try: short_groups_int = int(short_groups)
        except: short_groups_int = 2
    else:
        short_groups_int = 1

    # --- PAGE SETUP ---
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.header_distance = Inches(0.2)
    section.footer_distance = Inches(0.0)
    section.top_margin = Inches(0.3)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.3)
    section.right_margin = Inches(0.3)

    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(12)

    mcqs = ai_data.get("mcqs", [])
    short_qs = ai_data.get("short_qs", [])
    long_qs = ai_data.get("long_qs", [])

    mcq_marks = len(mcqs) * 1
    try: short_attempt_int = int(short_attempt)
    except: short_attempt_int = len(short_qs)
    short_marks_per_group = short_attempt_int * 2
    short_marks_total = short_marks_per_group * short_groups_int

    try: long_q_val = sum([int(m.strip()) for m in str(long_q_marks).split('+')])
    except: long_q_val = 9
    try: long_attempt_int = int(long_attempt)
    except: long_attempt_int = len(long_qs)
    long_marks = long_attempt_int * long_q_val
    total_marks = mcq_marks + short_marks_total + long_marks

    try:
        dt_obj = datetime.datetime.strptime(test_date, "%Y-%m-%d")
        formatted_date = dt_obj.strftime("%d %B %Y")
    except:
        formatted_date = test_date

    # ==========================================
    # CLEAN HEADER
    # ==========================================
    header = section.header
    for para in header.paragraphs:
        try:
            p = para._p
            p.getparent().remove(p)
        except: pass

    name_para = header.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_after = Pt(0)
    run_title = name_para.add_run(academy_name.upper())
    run_title.font.name = 'Times New Roman'
    run_title.font.size = Pt(26)
    run_title.font.bold = True

    p_sub = header.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(4)
    tabs = p_sub.paragraph_format.tab_stops
    tabs.add_tab_stop(Inches(3.83), WD_TAB_ALIGNMENT.CENTER)
    tabs.add_tab_stop(Inches(7.67), WD_TAB_ALIGNMENT.RIGHT)
    run_sub = p_sub.add_run(f"{subject}\t{class_name}\t{syllabus}")
    run_sub.font.name = 'Times New Roman'
    run_sub.font.size = Pt(12)
    run_sub.font.bold = True

    pPr = p_sub._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '18')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '305060')
    pBdr.append(bot)
    pPr.append(pBdr)

    # --- INFO ROW ---
    p_info = doc.add_paragraph()
    p_info.paragraph_format.space_after = Pt(8)
    info_tabs = p_info.paragraph_format.tab_stops
    info_tabs.add_tab_stop(Inches(2.6), WD_TAB_ALIGNMENT.LEFT)
    info_tabs.add_tab_stop(Inches(4.6), WD_TAB_ALIGNMENT.LEFT)
    info_tabs.add_tab_stop(Inches(7.67), WD_TAB_ALIGNMENT.RIGHT)

    r_name = p_info.add_run("Name: ")
    r_name.bold = True
    r_name.font.size = Pt(13)
    p_info.add_run("__________________\t")
    r_date = p_info.add_run("Date: ")
    r_date.bold = True
    r_date.font.size = Pt(13)
    p_info.add_run(f"{formatted_date}\t")
    r_time = p_info.add_run("Time: ")
    r_time.bold = True
    r_time.font.size = Pt(13)
    r_tval = p_info.add_run(f"{time_allowed}\t")
    r_tval.underline = True
    r_marks = p_info.add_run("Max Marks: ")
    r_marks.bold = True
    r_marks.font.size = Pt(13)
    r_mval = p_info.add_run(f" {total_marks} ")
    r_mval.underline = True

    # ==========================================
    # MCQs SECTION (Now uses exact left-right formatting)
    # ==========================================
    if mcqs:
        p_mcq_head = doc.add_paragraph()
        p_mcq_head.paragraph_format.space_after = Pt(6)
        run_l = p_mcq_head.add_run(f'Multiple Choice Questions ({mcq_marks} Marks)')
        run_l.bold = True
        run_l.font.size = Pt(14)

        if template_style == "column":
            # Best for bilingual!
            section_2 = doc.add_section(WD_SECTION_START.CONTINUOUS)
            sectPr2 = section_2._sectPr
            cols = sectPr2.xpath('./w:cols')[0]
            cols.set(qn('w:num'), '2')
            cols.set(qn('w:space'), '360')

            for idx, m in enumerate(mcqs, start=1):
                q_eng = m.get('q_en', m.get('question', ''))
                q_ur = m.get('q_ur', '')
                
                # Question Line
                if is_bilingual and q_ur:
                    add_side_by_side_bilingual(doc, clean_text(q_eng), clean_text(q_ur), num_prefix=f"{idx}. ", font_size=11, space_after=2, is_bold=True)
                else:
                    p_q = doc.add_paragraph()
                    p_q.paragraph_format.left_indent = Inches(0.2)
                    p_q.paragraph_format.first_line_indent = Inches(-0.2)
                    p_q.paragraph_format.space_after = Pt(2)
                    p_q.add_run(f"{idx}. ").bold = True
                    p_q.add_run(clean_text(q_eng)).bold = True

                # Options Lines
                opts = {
                    'A': (m.get('a_en', m.get('a','')), m.get('a_ur', '')),
                    'B': (m.get('b_en', m.get('b','')), m.get('b_ur', '')),
                    'C': (m.get('c_en', m.get('c','')), m.get('c_ur', '')),
                    'D': (m.get('d_en', m.get('d','')), m.get('d_ur', ''))
                }
                for key, (o_eng, o_ur) in opts.items():
                    if is_bilingual and o_ur:
                         # Use a slightly indented side-by-side for options
                         opt_tbl = add_side_by_side_bilingual(doc, clean_text(o_eng), clean_text(o_ur), num_prefix=f"    {key}) ", font_size=11, space_after=0)
                         opt_tbl.rows[0].cells[0].paragraphs[0].paragraph_format.left_indent = Inches(0.2)
                    else:
                        p_opt = doc.add_paragraph()
                        p_opt.paragraph_format.left_indent = Inches(0.4)
                        p_opt.paragraph_format.space_after = Pt(0)
                        p_opt.add_run(f"{key}) ").bold = True
                        p_opt.add_run(clean_text(o_eng))

                doc.paragraphs[-1].paragraph_format.space_after = Pt(8)

            section_3 = doc.add_section(WD_SECTION_START.CONTINUOUS)
            sectPr3 = section_3._sectPr
            cols3 = sectPr3.xpath('./w:cols')[0]
            cols3.set(qn('w:num'), '1')

        else:
            # Table Style (Fallback to 6-col table)
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            widths = [Inches(0.4), Inches(3.87), Inches(0.85), Inches(0.85), Inches(0.85), Inches(0.85)]
            for row in table.rows:
                for i, w in enumerate(widths): row.cells[i].width = w

            headers = ['No', 'Question', 'A', 'B', 'C', 'D']
            for i, txt in enumerate(headers):
                p = table.rows[0].cells[i].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run(txt).bold = True

            for idx, m in enumerate(mcqs, start=1):
                row_cells = table.add_row().cells
                for i, w in enumerate(widths):
                    row_cells[i].width = w
                    row_cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

                row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row_cells[0].paragraphs[0].add_run(str(idx)).bold = True
                
                # Q Text
                row_cells[1].paragraphs[0].add_run(clean_text(m.get('q_en', '')))
                if is_bilingual and m.get('q_ur'):
                    p_u = row_cells[1].add_paragraph()
                    set_rtl_paragraph(p_u)
                    add_urdu_run(p_u, clean_text(m.get('q_ur')), font_size=12)

                # Options
                opts_data = [('a_en','a_ur'), ('b_en','b_ur'), ('c_en','c_ur'), ('d_en','d_ur')]
                for cell_idx, (en_k, ur_k) in enumerate(opts_data, start=2):
                    row_cells[cell_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    row_cells[cell_idx].paragraphs[0].add_run(clean_text(m.get(en_k, '')))
                    if is_bilingual and m.get(ur_k):
                        p_ou = row_cells[cell_idx].add_paragraph()
                        set_rtl_paragraph(p_ou)
                        add_urdu_run(p_ou, clean_text(m.get(ur_k)), font_size=11)


    # ==========================================
    # SHORT QUESTIONS
    # ==========================================
    try: short_total_int = int(short_total)
    except: short_total_int = len(short_qs)

    for g in range(short_groups_int):
        group_start = g * short_total_int
        group_end = group_start + short_total_int
        group_qs = short_qs[group_start:group_end]
        if not group_qs: break

        q_number = g + 2
        sh = doc.add_paragraph()
        sh.paragraph_format.space_before = Pt(12)
        r_sq = sh.add_run(f'Q{q_number}. Short Questions (Attempt any {short_attempt} out of {short_total})')
        r_sq.bold = True
        r_sq.font.size = Pt(14)
        sh.paragraph_format.tab_stops.add_tab_stop(Inches(7.67), WD_TAB_ALIGNMENT.RIGHT)
        sh.add_run(f'\t{short_marks_per_group} Marks').bold = True

        for idx, sq in enumerate(group_qs, start=1):
            sq_eng = sq.get('text_en', sq.get('text', ''))
            sq_urdu = sq.get('text_ur', '')
            
            if is_bilingual and sq_urdu:
                add_side_by_side_bilingual(doc, clean_text(sq_eng), clean_text(sq_urdu), num_prefix=f"{idx}. ", font_size=12, space_after=6)
            else:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.first_line_indent = Inches(-0.25)
                p.paragraph_format.space_after = Pt(4)
                p.add_run(f"{idx}. ").bold = True
                p.add_run(clean_text(sq_eng))

    # ==========================================
    # LONG QUESTIONS
    # ==========================================
    if long_qs:
        lh = doc.add_paragraph()
        lh.paragraph_format.space_before = Pt(12)
        r_lq = lh.add_run(f'Long Questions (Attempt any {long_attempt} out of {long_total})')
        r_lq.bold = True
        r_lq.font.size = Pt(14)
        lh.paragraph_format.tab_stops.add_tab_stop(Inches(7.67), WD_TAB_ALIGNMENT.RIGHT)
        lh.add_run(f'\t{long_marks} Marks').bold = True

        for idx, lq in enumerate(long_qs, start=1):
            lq_eng = clean_text(lq.get('text_en', lq.get('text', ''))).replace('\\n', '\n')
            lq_urdu = clean_text(lq.get('text_ur', '')).replace('\\n', '\n')
            
            eng_parts = lq_eng.split('\n')
            urdu_parts = lq_urdu.split('\n') if lq_urdu else []

            for part_idx, part in enumerate(eng_parts):
                part = part.strip()
                if not part: continue

                u_part = urdu_parts[part_idx].strip() if part_idx < len(urdu_parts) else ""
                
                if is_bilingual and u_part:
                    prefix = f"{idx}. " if part_idx == 0 else "    "
                    add_side_by_side_bilingual(doc, part, u_part, num_prefix=prefix, font_size=12, space_after=6)
                else:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.25)
                    p.paragraph_format.first_line_indent = Inches(-0.25)
                    p.paragraph_format.space_after = Pt(4)
                    if part_idx == 0:
                        p.add_run(f"{idx}. ").bold = True
                    else:
                        p.paragraph_format.left_indent = Inches(0.5)
                        p.paragraph_format.first_line_indent = Inches(0)
                    p.add_run(part)

    output_path = "generated_test.docx"
    doc.save(output_path)
    return output_path